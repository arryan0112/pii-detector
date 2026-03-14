import csv
import io
import json
import re
import email as email_lib
from email import policy as email_policy
from pathlib import Path
from typing import Union

import chardet
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.ingestion.models import IngestionResult, TextChunk

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100


# ── Encoding ──────────────────────────────────────────────────────────────────

def _decode_bytes(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    detected = chardet.detect(raw)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        return raw.decode("utf-8", errors="replace")


# ── Chunking ──────────────────────────────────────────────────────────────────

def _chunk_text(text: str, source_file: str, fmt: str,
                location_prefix: str, start_chunk_idx: int = 0) -> list[TextChunk]:
    if len(text) <= CHUNK_SIZE:
        return [TextChunk(
            text=text,
            source_file=source_file,
            format=fmt,
            location=location_prefix,
            chunk_index=start_chunk_idx
        )]
    chunks = []
    pos = 0
    idx = start_chunk_idx
    while pos < len(text):
        end = min(pos + CHUNK_SIZE, len(text))
        chunks.append(TextChunk(
            text=text[pos:end],
            source_file=source_file,
            format=fmt,
            location=f"{location_prefix} (chars {pos}-{end})",
            chunk_index=idx
        ))
        pos += CHUNK_SIZE - CHUNK_OVERLAP
        idx += 1
    return chunks


# ── Extractors ────────────────────────────────────────────────────────────────

def _extract_plain_text(raw: bytes, filename: str) -> IngestionResult:
    text = _decode_bytes(raw)
    chunks = _chunk_text(text, filename, "txt", "full document")
    return IngestionResult(filename=filename, format="txt", chunks=chunks, errors=[])


def _extract_json(raw: bytes, filename: str) -> IngestionResult:
    text = _decode_bytes(raw)
    chunks = []
    errors = []

    def _walk(obj, path: str = "$"):
        if isinstance(obj, str):
            if obj.strip():
                chunks.append(TextChunk(
                    text=f"{path.split('.')[-1]}: {obj}",
                    source_file=filename,
                    format="json",
                    location=path
                ))
        elif isinstance(obj, dict):
            for k, v in obj.items():
                _walk(v, f"{path}.{k}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                _walk(item, f"{path}[{i}]")

    try:
        obj = json.loads(text)
        _walk(obj)
    except json.JSONDecodeError as e:
        errors.append(f"JSON parse error: {e}")
        chunks = _chunk_text(text, filename, "json", "raw text fallback")

    return IngestionResult(filename=filename, format="json", chunks=chunks, errors=errors)


def _extract_csv(raw: bytes, filename: str) -> IngestionResult:
    text = _decode_bytes(raw)
    chunks = []
    errors = []

    try:
        sample = text[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel

        reader = csv.DictReader(io.StringIO(text), dialect=dialect)

        for row_num, row in enumerate(reader, start=1):
            # Full row for cross-column context
            full_row = " | ".join(
                f"{k}: {v}" for k, v in row.items() if v and v.strip()
            )
            if full_row.strip():
                chunks.append(TextChunk(
                    text=full_row,
                    source_file=filename,
                    format="csv",
                    location=f"row {row_num} (full)",
                    chunk_index=row_num
                ))

            # Individual cells for precise location
            for col_name, value in row.items():
                if value and value.strip() and len(value.strip()) > 2:
                    chunks.append(TextChunk(
                        text=f"{col_name}: {value}",
                        source_file=filename,
                        format="csv",
                        location=f"row {row_num}, column '{col_name}'",
                        chunk_index=row_num,
                        metadata={"row": row_num, "column": col_name}
                    ))

    except Exception as e:
        errors.append(f"CSV parse error: {e}")
        chunks = _chunk_text(text, filename, "csv", "raw text fallback")

    return IngestionResult(filename=filename, format="csv", chunks=chunks, errors=errors)


def _extract_pdf(raw: bytes, filename: str) -> IngestionResult:
    chunks = []
    errors = []

    try:
        reader = PdfReader(io.BytesIO(raw))

        # Metadata
        meta = reader.metadata or {}
        meta_parts = []
        for key in ("/Author", "/Title", "/Subject", "/Keywords"):
            val = meta.get(key, "")
            if val and str(val).strip():
                meta_parts.append(f"{key.lstrip('/')}: {val}")
        if meta_parts:
            chunks.append(TextChunk(
                text="\n".join(meta_parts),
                source_file=filename,
                format="pdf",
                location="document metadata",
                chunk_index=0
            ))

        # Pages
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                text = (page.extract_text() or "").strip()
                if not text:
                    errors.append(f"Page {page_num}: no extractable text")
                    continue
                chunks.extend(_chunk_text(
                    text, filename, "pdf",
                    f"page {page_num}",
                    start_chunk_idx=page_num * 100
                ))
            except Exception as e:
                errors.append(f"Page {page_num} failed: {e}")

    except Exception as e:
        errors.append(f"PDF could not be opened: {e}")

    return IngestionResult(filename=filename, format="pdf", chunks=chunks, errors=errors)


def _extract_docx(raw: bytes, filename: str) -> IngestionResult:
    chunks = []
    errors = []

    try:
        doc = DocxDocument(io.BytesIO(raw))

        # Properties
        props = doc.core_properties
        prop_parts = []
        for attr in ("author", "last_modified_by", "title", "subject"):
            val = getattr(props, attr, None)
            if val:
                prop_parts.append(f"{attr}: {val}")
        if prop_parts:
            chunks.append(TextChunk(
                text="\n".join(prop_parts),
                source_file=filename,
                format="docx",
                location="document properties",
                chunk_index=0
            ))

        # Body
        body_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if body_text:
            chunks.extend(_chunk_text(body_text, filename, "docx", "body", start_chunk_idx=1))

        # Tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    chunks.append(TextChunk(
                        text=row_text,
                        source_file=filename,
                        format="docx",
                        location=f"table {t_idx + 1}, row {r_idx + 1}",
                        chunk_index=1000 + t_idx * 100 + r_idx
                    ))

    except Exception as e:
        errors.append(f"DOCX could not be opened: {e}")

    return IngestionResult(filename=filename, format="docx", chunks=chunks, errors=errors)


def _strip_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    parts = []
    if soup.title and soup.title.string:
        parts.append(f"Page title: {soup.title.string.strip()}")
    for meta in soup.find_all("meta"):
        name = meta.get("name", "").lower()
        content = meta.get("content", "").strip()
        if name in ("description", "keywords", "author") and content:
            parts.append(f"{name}: {content}")
    for img in soup.find_all("img", alt=True):
        if img["alt"].strip():
            parts.append(f"image alt: {img['alt'].strip()}")

    visible = soup.get_text(separator="\n")
    visible = re.sub(r"\n{3,}", "\n\n", visible)
    visible = re.sub(r"[ \t]{2,}", " ", visible)
    parts.append(visible.strip())
    return "\n\n".join(filter(None, parts))


def _extract_html(raw: bytes, filename: str) -> IngestionResult:
    text = _decode_bytes(raw)
    extracted = _strip_html(text)
    chunks = _chunk_text(extracted, filename, "html", "page content")
    return IngestionResult(filename=filename, format="html", chunks=chunks, errors=[])


def _extract_email(raw: bytes, filename: str) -> IngestionResult:
    chunks = []
    errors = []

    try:
        msg = email_lib.message_from_bytes(raw, policy=email_policy.default)

        # Headers
        header_parts = []
        for header in ("From", "To", "CC", "BCC", "Subject"):
            val = msg.get(header, "")
            if val:
                header_parts.append(f"{header}: {val}")
        if header_parts:
            chunks.append(TextChunk(
                text="\n".join(header_parts),
                source_file=filename,
                format="email",
                location="email headers"
            ))

        # Body parts
        for part in msg.walk():
            content_type = part.get_content_type()
            disposition = part.get_content_disposition() or ""

            if disposition == "attachment":
                attachment_filename = part.get_filename() or "attachment"
                try:
                    attachment_bytes = part.get_payload(decode=True)
                    if attachment_bytes:
                        sub_result = ingest_bytes(attachment_bytes, attachment_filename)
                        for chunk in sub_result.chunks:
                            chunk.location = f"attachment '{attachment_filename}' -> {chunk.location}"
                            chunk.source_file = filename
                            chunks.append(chunk)
                except Exception as e:
                    errors.append(f"Attachment '{attachment_filename}' failed: {e}")

            elif content_type == "text/plain":
                body = part.get_payload(decode=True)
                if body:
                    text = _decode_bytes(body)
                    chunks.extend(_chunk_text(text, filename, "email", "email body"))

            elif content_type == "text/html":
                body = part.get_payload(decode=True)
                if body:
                    html = _decode_bytes(body)
                    text = _strip_html(html)
                    if text.strip():
                        chunks.extend(_chunk_text(text, filename, "email", "email body (html)"))

    except Exception as e:
        errors.append(f"Email parse error: {e}")

    return IngestionResult(filename=filename, format="email", chunks=chunks, errors=errors)


def _extract_log(raw: bytes, filename: str) -> IngestionResult:
    text = _decode_bytes(raw)
    lines = text.splitlines()
    chunks = []
    BATCH = 50

    for i in range(0, len(lines), BATCH):
        batch = lines[i:i + BATCH]
        batch_text = "\n".join(line for line in batch if line.strip())
        if batch_text:
            chunks.append(TextChunk(
                text=batch_text,
                source_file=filename,
                format="log",
                location=f"lines {i + 1}-{i + len(batch)}",
                chunk_index=i // BATCH
            ))

    return IngestionResult(filename=filename, format="log", chunks=chunks, errors=[])


# ── Format detection ──────────────────────────────────────────────────────────

MAGIC_BYTES = {
    b"%PDF":              "pdf",
    b"PK\x03\x04":       "docx",
    b"\xd0\xcf\x11\xe0": "doc",
}

EXTENSION_MAP = {
    ".txt": "txt", ".log": "log", ".csv": "csv",
    ".json": "json", ".jsonl": "json",
    ".pdf": "pdf", ".docx": "docx",
    ".htm": "html", ".html": "html",
    ".eml": "email",
}

EXTRACTORS = {
    "txt":   _extract_plain_text,
    "log":   _extract_log,
    "csv":   _extract_csv,
    "json":  _extract_json,
    "pdf":   _extract_pdf,
    "docx":  _extract_docx,
    "email": _extract_email,
    "html":  _extract_html,
}


def _detect_format(filename: str, raw: bytes) -> str:
    for magic, fmt in MAGIC_BYTES.items():
        if raw.startswith(magic):
            ext = Path(filename).suffix.lower()
            if fmt == "docx" and ext == ".xlsx":
                return "xlsx"
            return fmt
    ext = Path(filename).suffix.lower()
    if ext in EXTENSION_MAP:
        return EXTENSION_MAP[ext]
    sample = raw[:512].strip()
    if sample.startswith(b"{") or sample.startswith(b"["):
        return "json"
    return "txt"


# ── Public API ────────────────────────────────────────────────────────────────

def ingest_bytes(raw: bytes, filename: str) -> IngestionResult:
    fmt = _detect_format(filename, raw)
    extractor = EXTRACTORS.get(fmt)

    if extractor is None:
        return IngestionResult(
            filename=filename, format=fmt, chunks=[],
            errors=[f"Unsupported format: {fmt}"]
        )

    try:
        return extractor(raw, filename)
    except Exception as e:
        return IngestionResult(
            filename=filename, format=fmt, chunks=[],
            errors=[f"Fatal ingestion error: {e}"]
        )


def ingest_file(path: Union[str, Path]) -> IngestionResult:
    path = Path(path)
    if not path.exists():
        return IngestionResult(
            filename=str(path), format="unknown",
            chunks=[], errors=[f"File not found: {path}"]
        )
    return ingest_bytes(path.read_bytes(), path.name)


def ingest_text(text: str, source: str = "direct_input") -> IngestionResult:
    chunks = _chunk_text(text, source, "text", "direct input")
    return IngestionResult(
        filename=source, format="text", chunks=chunks, errors=[]
    )